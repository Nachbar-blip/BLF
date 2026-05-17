"""Generiert BLF-Mathe_QR.docx mit QR-Code zur Lernapp."""
from pathlib import Path
import qrcode
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

URL = "https://nachbar-blip.github.io/BLF/"
OUT_DOCX = Path(__file__).parent / "docs" / "BLF-Mathe_QR.docx"
OUT_PNG = Path(__file__).parent / "docs" / "qr.png"


def make_qr():
    qr = qrcode.QRCode(version=None, error_correction=qrcode.constants.ERROR_CORRECT_M, box_size=12, border=4)
    qr.add_data(URL)
    qr.make(fit=True)
    img = qr.make_image(fill_color="#1e40af", back_color="white")
    img.save(OUT_PNG)


def make_docx():
    doc = Document()

    # Seitenränder etwas enger
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # Titel
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("BLF Mathe — Thüringen")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)

    # Untertitel
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.add_run("Vorbereitung auf die Besondere Leistungsfeststellung — Klasse 10")
    sub_run.italic = True
    sub_run.font.size = Pt(13)
    sub_run.font.color.rgb = RGBColor(0x4B, 0x5B, 0x6B)

    doc.add_paragraph()

    # QR-Code zentriert einfügen
    qr_p = doc.add_paragraph()
    qr_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    qr_p.add_run().add_picture(str(OUT_PNG), width=Cm(8))

    # Link unter dem QR-Code
    link_p = doc.add_paragraph()
    link_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    link_run = link_p.add_run(URL)
    link_run.font.size = Pt(11)
    link_run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)

    doc.add_paragraph()

    # Anleitung — kurz
    intro = doc.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.LEFT
    intro_run = intro.add_run("So nutzt du die App")
    intro_run.bold = True
    intro_run.font.size = Pt(14)

    for line in [
        "Du wählst zwischen 8 BLF-Themen (Arithmetik, Funktionen, Geometrie, Stochastik).",
        "Pro Thema gibt es zwei Modi:",
    ]:
        p = doc.add_paragraph(line)
        p.paragraph_format.space_after = Pt(2)

    bp = doc.add_paragraph(style="List Bullet")
    bp.add_run("📚 Lernen (Karteikarten): ").bold = True
    bp.add_run("Begriffe, Formeln und Sätze einprägen — Leitner-Box mit 5 Stufen.")

    bp = doc.add_paragraph(style="List Bullet")
    bp.add_run("🧠 Üben (Aufgaben): ").bold = True
    bp.add_run("adaptive Aufgaben (Level 1–6), die sich automatisch deinem Niveau anpassen.")

    p = doc.add_paragraph("Aufgaben mit dem orangefarbenen Badge ")
    p.add_run("hilfsmittelfrei").bold = True
    p.add_run(" sind für Teil A der BLF (ohne Taschenrechner).")
    p.paragraph_format.space_before = Pt(8)

    p = doc.add_paragraph()
    p.add_run("Dein Fortschritt wird automatisch im Browser gespeichert — kein Login nötig.").italic = True

    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_DOCX)
    print(f"Geschrieben: {OUT_DOCX}")


if __name__ == "__main__":
    make_qr()
    make_docx()
